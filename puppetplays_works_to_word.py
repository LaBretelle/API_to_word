import re
import json
import requests as requests
from docx import Document
import os.path
import textract


directory = "liste_oeuvres"
if not os.path.exists(directory):
    os.makedirs(directory)



#firstly, fetch data using the PuppetPlays API
url = 'https://api.puppetplays.eu/graphql/'
query = """{ entries(typeId:15) { title note mainTheme abstract }}""" #{ ... on works_works_Entry authors { title } }
r = requests.post(url, json={'query': query})

#load data and transform it
json_data = json.loads(r.text)
liste_oeuvres = json_data['data']['entries']

for oeuvre in liste_oeuvres:
    #print(type(oeuvre), oeuvre)
    document = Document() #init word document to write


    #Écriture du titre de la page courante
    document.add_heading(oeuvre['mainTheme'], 0)

    document.add_heading('Notice', level=1)

    #Retrait des balises pour une lecture humaine.
    pattern = '<[^>]*>'
    replace = ''
    notice = re.sub(pattern, replace, oeuvre['note'])
    resume = re.sub(pattern, replace, oeuvre['abstract'])

    #Ecris la notice.
    p = document.add_paragraph('\n')
    p.add_run(notice)

    #Écriture du résumé.
    document.add_heading('Résumé', level=1)
    p = document.add_paragraph('\n')
    p.add_run(resume)

    #Ajout du titre et suppression des éléments causant des erreurs dans les noms de fichiers
    titre = oeuvre['title']
    titre = titre.replace(" ", "_")
    titre = titre.replace("\'", "_")
    titre = titre.replace("?", "")
    titre = titre.replace(",", "")
    titre = titre.replace("-", "_")

    #initialise le nom du fichier
    creation = ".\\" + directory + "\\" + titre + ".docx"

    # permet de ne pas écraser les oeuvres ayant le même titre en leur ajoutant un numéro
    counter = 1
    while os.path.exists(creation):
        if resume in textract.process(creation).decode('utf-8'):
            break
        print(titre)
        counter += 1
        creation = ".\\" + directory + "\\" + titre + "_" + str(counter) + ".docx"
        print(counter)
    document.save(creation)
